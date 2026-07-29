# -*- coding: utf-8 -*-
"""
מודול הרכבה משותף לקובצי וורד בעברית מיושרים לימין.

כלל היישום (הכלל הקבוע מכאן ולהבא):
פסקה בעברית מסומנת RTL באמצעות w:bidi בלבד, ואיננה מקבלת w:jc="right".
הסיבה: בפסקה מסוג RTL, מנועי תצוגה מודרניים (גוגל דוקס, מציגים ניידים,
תצוגות מבוססות ובכללן הפרשנות ה"קפדנית" של OOXML) קוראים את הערך "right"
באופן לוגי כ"end", וקצה השורה בכתיבה מימין לשמאל הוא הצד השמאלי, ולכן
הטקסט נצמד לשמאל. פסקת bidi ללא jc נצמדת מאליה לתחילת השורה, שהיא הצד הימני,
וכל המנועים מסכימים על כך. jc מפורש נשמור רק למרכוז (center), שאין בו
עמימות ימין/שמאל.
"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# אלמנטים שמותר להם לבוא אחרי w:bidi בתוך w:pPr, לפי סדר הסכמה של OOXML
_BIDI_SUCCESSORS = (
    'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
    'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
    'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
    'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange',
)


def insert_bidi(pPr):
    b = pPr.find(qn('w:bidi'))
    if b is None:
        b = OxmlElement('w:bidi')
        pPr.insert_element_before(b, *_BIDI_SUCCESSORS)
    b.set(qn('w:val'), '1')
    return b


def _remove_jc(pPr):
    jc = pPr.find(qn('w:jc'))
    if jc is not None:
        pPr.remove(jc)


def set_rtl_style(style, size=None, bold=None):
    style.font.name = "David"
    r = style.element.get_or_add_rPr()
    rf = r.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); r.insert(0, rf)
    rf.set(qn('w:ascii'), 'David'); rf.set(qn('w:hAnsi'), 'David'); rf.set(qn('w:cs'), 'David')
    if size:
        style.font.size = Pt(size)
        szCs = r.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs'); r.append(szCs)
        szCs.set(qn('w:val'), str(int(size * 2)))
    if bold:
        style.font.bold = True
        bCs = r.find(qn('w:bCs'))
        if bCs is None:
            bCs = OxmlElement('w:bCs'); r.append(bCs)
        bCs.set(qn('w:val'), '1')
    try:
        pPr = style.element.get_or_add_pPr()
        insert_bidi(pPr)
        _remove_jc(pPr)  # יישור טבעי לימין בירושה, בלי jc עמים
    except AttributeError:
        pass


def apply_base_styles(doc, headings):
    """headings: רשימת (שם סגנון, גודל) להחלת גופן David ו-RTL."""
    set_rtl_style(doc.styles['Normal'], size=12)
    for h, s in headings:
        try:
            set_rtl_style(doc.styles[h], size=s, bold=(h != 'List Bullet'))
        except KeyError:
            pass


# ברירת המחדל ליישור פסקאות הגוף. 'both' = יישור דו-צדדי (justify), לבקשת
# המחבר; אושר כתקין בוורד בנייד (הג'יבריש שנראה תחילה היה עיכוב טעינת גופן
# חולף בלבד). 'right' = יישור ימין טבעי בלי jc, חלופה זמינה.
BODY_ALIGN = 'both'


def make_rtl(p, center=False, align=None):
    """מסמן פסקה כ-RTL.
    center=True  -> מרכוז (jc=center), לכותרות.
    align='both' -> יישור דו-צדדי (justify).
    align='right'/None -> יישור ימין טבעי, בלי jc.
    אם align לא נמסר, נגזר מ-BODY_ALIGN."""
    pPr = p._p.get_or_add_pPr()
    insert_bidi(pPr)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # מרכוז סימטרי, ללא עמימות
    else:
        eff = align or BODY_ALIGN
        if eff == 'both':
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # jc=both, חד-משמעי
        else:
            _remove_jc(pPr)                           # יישור ימין טבעי של RTL
    for run in p.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = rPr.find(qn('w:rtl'))
        if rtl is None:
            rtl = OxmlElement('w:rtl'); rPr.append(rtl)
        elif rtl.get(qn('w:val')) in ('0', 'false'):
            continue  # ריצה שסומנה במפורש LTR (למשל ספרות מספור) נשמרת כפי שהיא
        rtl.set(qn('w:val'), '1')
    return p


def _set_run_rtl(run, rtl=True):
    rPr = run._r.get_or_add_rPr()
    el = rPr.find(qn('w:rtl'))
    if el is None:
        el = OxmlElement('w:rtl'); rPr.append(el)
    el.set(qn('w:val'), '1' if rtl else '0')
    return run


def add_clause_title(doc, text, size=None):
    """כותרת סעיף: מודגשת ובלא מספר. המספר שייך לגוף הסעיף, לא לכותרת."""
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = True
    if size:
        run.font.size = Pt(size)
    bCs = OxmlElement('w:bCs'); bCs.set(qn('w:val'), '1')
    run._r.get_or_add_rPr().append(bCs)
    make_rtl(p, align='right')
    return p


def add_numbered_clause(doc, num, text, bold=False):
    """סעיף ממוספר, לפי כלל המספור הקבוע.

    המספר בא לפני גוף הסעיף (ולא על הכותרת), והנקודה באה אחרי המספר, "11."
    ולא ".11". לשם כך הספרות יושבות בריצת LTR נפרדת, והנקודה והרווחים בריצת
    RTL הבאה אחריה. נקודה שנשארת בתוך ריצת הספרות נסחפת ומוצגת לפני המספר.
    """
    p = doc.add_paragraph()
    _set_run_rtl(p.add_run(str(num)), rtl=False)   # ספרות בלבד, LTR
    _set_run_rtl(p.add_run('.  '), rtl=True)  # נקודה ורווחים, RTL
    body = p.add_run(text)
    if bold:
        body.bold = True
        bCs = OxmlElement('w:bCs'); bCs.set(qn('w:val'), '1')
        body._r.get_or_add_rPr().append(bCs)
    _set_run_rtl(body, rtl=True)
    make_rtl(p)
    return p


def add_letter_item(doc, letter, text):
    """תת סעיף באות עברית. הסוגריים, האות והרווחים כולם בריצת RTL אחת."""
    p = doc.add_paragraph()
    _set_run_rtl(p.add_run('(%s)  ' % letter), rtl=True)
    _set_run_rtl(p.add_run(text), rtl=True)
    make_rtl(p)
    return p


def set_section_rtl(section):
    sectPr = section._sectPr
    if sectPr.find(qn('w:bidi')) is None:
        sectPr.append(OxmlElement('w:bidi'))


def set_doc_defaults_rtl(doc):
    """מוסיף w:bidi ו-w:rtl לברירת המחדל של המסמך (docDefaults),
    כדי שגם וורד בנייד, הקורא את ברירות המחדל תחילה, יתחיל מ-RTL."""
    styles = doc.styles.element  # w:styles
    dd = styles.find(qn('w:docDefaults'))
    if dd is None:
        dd = OxmlElement('w:docDefaults'); styles.insert(0, dd)
    # ברירת מחדל לפסקה
    pdd = dd.find(qn('w:pPrDefault'))
    if pdd is None:
        pdd = OxmlElement('w:pPrDefault'); dd.append(pdd)
    ppr = pdd.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr'); pdd.append(ppr)
    if ppr.find(qn('w:bidi')) is None:
        ppr.insert(0, OxmlElement('w:bidi'))
    # ברירת מחדל לריצה
    rdd = dd.find(qn('w:rPrDefault'))
    if rdd is None:
        rdd = OxmlElement('w:rPrDefault'); dd.append(rdd)
    rpr = rdd.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr'); rdd.append(rpr)
    if rpr.find(qn('w:rtl')) is None:
        rpr.append(OxmlElement('w:rtl'))


def add_page_number_footer(section):
    fp = section.footer.paragraphs[0]
    insert_bidi(fp._p.get_or_add_pPr())
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), ' PAGE \\* MERGEFORMAT ')
    fr = OxmlElement('w:r')
    frPr = OxmlElement('w:rPr')
    frf = OxmlElement('w:rFonts')
    frf.set(qn('w:ascii'), 'David'); frf.set(qn('w:hAnsi'), 'David'); frf.set(qn('w:cs'), 'David')
    frPr.append(frf); fr.append(frPr)
    ft = OxmlElement('w:t'); ft.text = '1'; fr.append(ft)
    fld.append(fr); fp._p.append(fld)


def qa_report(path):
    """בקרה: כל פסקת תוכן חייבת bidi; ואסור jc=right/left בפסקת RTL."""
    from docx import Document
    d = Document(path)
    content = [p for p in d.paragraphs if p.text.strip()]
    no_bidi = bad_jc = 0
    for p in content:
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None or pPr.find(qn('w:bidi')) is None:
            no_bidi += 1; continue
        jc = pPr.find(qn('w:jc'))
        if jc is not None and jc.get(qn('w:val')) in ('right', 'left', 'start', 'end'):
            bad_jc += 1  # ערך יישור עמים בפסקת RTL
    foot_ok = bool(d.sections) and 'PAGE' in d.sections[0].footer.paragraphs[0]._p.xml
    return {'paragraphs': len(content), 'no_bidi': no_bidi,
            'ambiguous_jc': bad_jc, 'page_field': foot_ok}
