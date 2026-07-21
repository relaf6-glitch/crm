# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/home/user/crm/עבודה-חוברת"
OUT = BASE + "/תוצר/תקציר-מנהלים.docx"
SRC = BASE + "/תקציר/תקציר-מנהלים.md"

BIDI_SUCCESSORS = ('w:adjustRightInd','w:snapToGrid','w:spacing','w:ind','w:contextualSpacing','w:mirrorIndents','w:suppressOverlap','w:jc','w:textDirection','w:textAlignment','w:textboxTightWrap','w:outlineLvl','w:divId','w:cnfStyle','w:rPr','w:sectPr','w:pPrChange')

def insert_bidi(pPr):
    b = pPr.find(qn('w:bidi'))
    if b is None:
        b = OxmlElement('w:bidi'); pPr.insert_element_before(b, *BIDI_SUCCESSORS)
    b.set(qn('w:val'), '1')

doc = Document()

def set_rtl_style(style, size=None, bold=None):
    style.font.name = "David"
    r = style.element.get_or_add_rPr()
    rf = r.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); r.insert(0, rf)
    rf.set(qn('w:ascii'),'David'); rf.set(qn('w:hAnsi'),'David'); rf.set(qn('w:cs'),'David')
    if size:
        style.font.size = Pt(size)
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(int(size*2))); r.append(szCs)
    if bold:
        style.font.bold = True
        bCs = OxmlElement('w:bCs'); bCs.set(qn('w:val'),'1'); r.append(bCs)
    try: insert_bidi(style.element.get_or_add_pPr())
    except AttributeError: pass

set_rtl_style(doc.styles['Normal'], size=12)
for h,s in [('Heading 1',20),('Heading 2',16)]:
    set_rtl_style(doc.styles[h], size=s, bold=True)

def make_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p.alignment = align
    insert_bidi(p._p.get_or_add_pPr())
    for run in p.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'),'1'); rPr.append(rtl)

def add_runs(p, text):
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); run.bold = True
            bCs = OxmlElement('w:bCs'); bCs.set(qn('w:val'),'1'); run._r.get_or_add_rPr().append(bCs)
        elif part:
            p.add_run(part)

def add_paragraph(text, style=None):
    p = doc.add_paragraph(style=style); add_runs(p, text); make_rtl(p); return p

section = doc.sections[0]
sec_bidi = OxmlElement('w:bidi'); section._sectPr.append(sec_bidi)
fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
insert_bidi(fp._p.get_or_add_pPr())
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),' PAGE \\* MERGEFORMAT ')
fr = OxmlElement('w:r'); ft = OxmlElement('w:t'); ft.text='1'; fr.append(ft); fld.append(fr); fp._p.append(fld)

for line in open(SRC, encoding='utf-8').read().splitlines():
    s = line.rstrip()
    if not s.strip() or s.strip()=='---': continue
    if s.startswith('## '): add_paragraph(s[3:], style='Heading 2')
    elif s.startswith('# '): add_paragraph(s[2:], style='Heading 1')
    else: add_paragraph(s)

doc.save(OUT)
print('נשמר:', OUT)
from docx import Document as D2
d = D2(OUT); content=[p for p in d.paragraphs if p.text.strip()]; bad=0
for p in content:
    pPr=p._p.find(qn('w:pPr'))
    ok = pPr is not None and pPr.find(qn('w:bidi')) is not None and p.alignment in (WD_ALIGN_PARAGRAPH.RIGHT,WD_ALIGN_PARAGRAPH.CENTER)
    if not ok: bad+=1
foot = 'PAGE' in d.sections[0].footer.paragraphs[0]._p.xml
print('פסקאות:',len(content),'| חריגות יישור:',bad,'| מספור עמודים:', 'קיים' if foot else 'חסר')
