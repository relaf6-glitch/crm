# -*- coding: utf-8 -*-
# הרכבת החוברת המלאה לקובץ וורד מיושר לימין (RTL טבעי, גופן David, מספור עמודים)
import re, glob, sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from rtl_docx import (apply_base_styles, make_rtl, insert_bidi,
                      set_section_rtl, set_doc_defaults_rtl, add_page_number_footer, qa_report)

BASE = "/home/user/crm/עבודה-חוברת"
OUT = BASE + "/תוצר/חוברת-מדינה-יהודית-בעלת-תוכן.docx"
FILES = [BASE + "/מבוא/מבוא-מעודכן.md"] + sorted(glob.glob(BASE + "/פרקים/*.md")) + [BASE + "/נספחים.md"]

doc = Document()
apply_base_styles(doc, [('Heading 1', 20), ('Heading 2', 16), ('Heading 3', 14),
                        ('Title', 26), ('List Bullet', 12)])
set_doc_defaults_rtl(doc)
set_section_rtl(doc.sections[0])
add_page_number_footer(doc.sections[0])


def add_runs(p, text):
    text = re.sub(r'\[\^([^\]]+)\]', r'[\1]', text)
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); run.bold = True
            bCs = OxmlElement('w:bCs'); bCs.set(qn('w:val'), '1'); run._r.get_or_add_rPr().append(bCs)
        elif part:
            p.add_run(part)


def add_paragraph(text, style=None, center=False):
    p = doc.add_paragraph(style=style); add_runs(p, text); make_rtl(p, center=center); return p


# עמוד שער
t = doc.add_paragraph(style='Title'); add_runs(t, 'מדינה יהודית בעלת תוכן'); make_rtl(t, center=True)
st = doc.add_paragraph(); add_runs(st, 'עשרה מאמרים על זהות, ריבונות ומשפט'); make_rtl(st, center=True)
st.runs[0].font.size = Pt(16)
doc.add_page_break()

# תוכן עניינים ידני
toc_h = doc.add_paragraph(style='Heading 1'); add_runs(toc_h, 'תוכן העניינים'); make_rtl(toc_h)
toc_items = ['מבוא']
for f in FILES[1:]:
    first = open(f, encoding='utf-8').readline().strip()
    if first.startswith('#'):
        toc_items.append(first.lstrip('#').strip())
for item in toc_items:
    add_paragraph(item)

for f in FILES:
    for line in open(f, encoding='utf-8').read().splitlines():
        s = line.rstrip()
        if not s.strip() or s.strip() == '---':
            continue
        if s.startswith('### '):
            add_paragraph(s[4:], style='Heading 3')
        elif s.startswith('## '):
            add_paragraph(s[3:], style='Heading 2')
        elif s.startswith('# '):
            doc.add_page_break()
            add_paragraph(s[2:], style='Heading 1')
        elif s.startswith('- '):
            add_paragraph(s[2:], style='List Bullet')
        elif re.match(r'\[\^[^\]]+\]:', s):
            p = doc.add_paragraph()
            add_runs(p, re.sub(r'^\[\^([^\]]+)\]:', r'[\1]', s))
            for run in p.runs:
                run.font.size = Pt(9)
                szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '18'); run._r.get_or_add_rPr().append(szCs)
            make_rtl(p)
        else:
            add_paragraph(s)

doc.save(OUT)
print('נשמר:', OUT)
print('בקרה:', qa_report(OUT))
