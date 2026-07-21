# -*- coding: utf-8 -*-
# קובץ וורד: המבוא המעודכן + ארבע הצעות השינוי, לעיון המחבר (RTL טבעי)
import re, sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from rtl_docx import apply_base_styles, make_rtl, set_section_rtl, set_doc_defaults_rtl, add_page_number_footer, qa_report

BASE = "/home/user/crm/עבודה-חוברת"
OUT = BASE + "/תוצר/מבוא-והצעות-שינוי-לאישור.docx"
FILES = [BASE + "/מבוא/מבוא-מעודכן.md", BASE + "/מבוא/שינויי-מבוא.md"]

doc = Document()
apply_base_styles(doc, [('Heading 1', 20), ('Heading 2', 16), ('Heading 3', 14), ('List Bullet', 12)])
set_doc_defaults_rtl(doc)
set_section_rtl(doc.sections[0])
add_page_number_footer(doc.sections[0])


def add_runs(p, text):
    text = re.sub(r'\[\^([^\]]+)\]', r'[\1]', text)
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); run.bold = True
            bCs = OxmlElement('w:bCs'); bCs.set(qn('w:val'), '1'); run._r.get_or_add_rPr().append(bCs)
        elif part:
            p.add_run(part)


def add_paragraph(text, style=None):
    p = doc.add_paragraph(style=style); add_runs(p, text); make_rtl(p); return p


first = True
for f in FILES:
    if not first:
        doc.add_page_break()
    first = False
    for line in open(f, encoding='utf-8').read().splitlines():
        s = line.rstrip()
        if not s.strip() or s.strip() == '---':
            continue
        if s.startswith('### '):
            add_paragraph(s[4:], style='Heading 3')
        elif s.startswith('## '):
            add_paragraph(s[3:], style='Heading 2')
        elif s.startswith('# '):
            add_paragraph(s[2:], style='Heading 1')
        elif s.startswith('- '):
            add_paragraph(s[2:], style='List Bullet')
        else:
            add_paragraph(s)

doc.save(OUT)
print('נשמר:', OUT)
print('בקרה:', qa_report(OUT))
